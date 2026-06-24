from pathlib import Path
from abc import ABC, abstractmethod

import mistletoe
from mistletoe.base_renderer import BaseRenderer
import xml.etree.ElementTree as ElementTree


class BaseReader(ABC):
    @abstractmethod
    def read(self, path: Path) -> str:
        pass


class PlainTextReader(BaseReader):
    def read(self, path: Path) -> str:
        return path.read_text(encoding="utf-8")


class DocxReader(BaseReader):
    def read(self, path: Path) -> str:
        import docx

        doc = docx.Document(path)
        text_runs = []
        for p in doc.paragraphs:
            if p.text:
                text_runs.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_runs.append(cell.text)
        return "\n".join(text_runs)


class XliffReader(BaseReader):
    def read(self, path: Path) -> str:
        tree = ElementTree.parse(path)
        root = tree.getroot()
        ns_uri = ""
        if root.tag.startswith("{"):
            ns_uri = root.tag.split("}")[0].strip("{")

        version = root.attrib.get("version", "1.2")
        if version.startswith("2."):
            return self._parse_xliff_2(root, ns_uri)
        else:
            return self._parse_xliff_1(root, ns_uri)

    def _parse_xliff_1(self, root: ElementTree.Element, ns_uri: str) -> str:
        ns = {"ns": ns_uri} if ns_uri else {}
        xpath = ".//ns:source" if ns_uri else ".//source"
        texts = []
        for elem in root.findall(xpath, ns):
            if elem.text:
                texts.append(elem.text)
        return "\n".join(texts)

    def _parse_xliff_2(self, root: ElementTree.Element, ns_uri: str) -> str:
        ns = {"ns": ns_uri} if ns_uri else {}
        xpath = ".//ns:source" if ns_uri else ".//source"
        texts = []
        for elem in root.findall(xpath, ns):
            if elem.text:
                texts.append(elem.text)
        return "\n".join(texts)


class PlainTextRenderer(BaseRenderer):
    def __init__(self):
        super().__init__()

    def render_inner(self, token) -> str:
        if hasattr(token, "children") and token.children:
            return "".join(self.render(child) for child in token.children)
        return ""

    def render_line_break(self, token) -> str:
        return "\n"

    def render_raw_text(self, token) -> str:
        return token.content

    def render_strong(self, token) -> str:
        return self.render_inner(token)

    def render_emphasis(self, token) -> str:
        return self.render_inner(token)

    def render_inline_code(self, token) -> str:
        return token.children[0].content

    def render_link(self, token) -> str:
        return self.render_inner(token)

    def render_paragraph(self, token) -> str:
        return self.render_inner(token) + "\n\n"

    def render_heading(self, token) -> str:
        return self.render_inner(token) + "\n"

    def render_list_item(self, token) -> str:
        return self.render_inner(token)

    def render_block_code(self, token) -> str:
        return ""


class MarkdownReader(BaseReader):
    def read(self, path: Path) -> str:
        doc = mistletoe.Document(path.read_text(encoding="utf-8"))
        with PlainTextRenderer() as renderer:
            return renderer.render(doc)


class ReaderFactory:
    _readers = {
        ".txt": PlainTextReader(),
        ".docx": DocxReader(),
        ".xlf": XliffReader(),
        ".xliff": XliffReader(),
        ".md": MarkdownReader(),
        ".markdown": MarkdownReader(),
    }

    @classmethod
    def get_reader(cls, path: Path) -> BaseReader:
        return cls._readers.get(path.suffix.lower(), PlainTextReader())
